#!/usr/bin/env python3
"""
Convert ARCHITECTURE_OVERVIEW.md to a Word document (.docx) with embedded Mermaid diagram images.

Uses python-docx proper API for formatting (not HTML injection).

Usage:
    pip install python-docx
    python convert_md_to_docx.py

Requires: @mermaid-js/mermaid-cli installed globally (npm install -g @mermaid-js/mermaid-cli)
"""

import re
import subprocess
import sys
from pathlib import Path

# python-docx imports
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_mermaid_blocks(markdown_text: str) -> list[dict]:
    """Extract all mermaid diagram blocks from markdown text."""
    results = []
    for m in re.finditer(r'```mermaid\n(.*?)\n```', markdown_text, re.DOTALL):
        results.append({'index': len(results), 'content': m.group(1).strip()})
    return results


def render_mermaid_to_png(diagram_content: str, output_path: Path, width: int = 1200) -> bool:
    """Render a mermaid diagram to PNG using mmdc."""
    try:
        tmp_file = output_path.with_suffix('.mmd')
        tmp_file.write_text(diagram_content, encoding='utf-8')

        cmd = ['mmdc.cmd', '-i', str(tmp_file), '-o', str(output_path), '-w', str(width)]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, cwd=output_path.parent)

        if tmp_file.exists():
            tmp_file.unlink()

        return result.returncode == 0 and output_path.exists()
    except Exception as e:
        print(f"Error rendering mermaid diagram: {e}", file=sys.stderr)
        return False


def render_all_diagrams(diagram_dir: Path, diagrams: list[dict]) -> dict[int, Path]:
    """Render all Mermaid diagrams to PNG images. Returns index -> image_path mapping."""
    rendered = {}
    for i, diagram in enumerate(diagrams):
        img_path = diagram_dir / f"diagram_{i}.png"
        print(f"Rendering diagram {i}...")
        if render_mermaid_to_png(diagram['content'], img_path):
            rendered[i] = img_path
            print(f"  -> Saved to {img_path}")
        else:
            print(f"  -> FAILED", file=sys.stderr)
    return rendered


def parse_markdown_into_elements(markdown_text: str, diagrams: list[dict]) -> list[dict]:
    """Parse markdown text into structured elements for python-docx."""
    lines = markdown_text.split('\n')
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Check if this is a mermaid block start (```mermaid)
        if stripped.startswith('```mermaid'):
            # Find the end of the mermaid block
            content_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                content_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            elements.append({
                'type': 'mermaid',
                'content': '\n'.join(content_lines).strip()
            })
            continue

        # Code blocks (skip - they're just text)
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append({'type': 'code', 'language': lang, 'content': '\n'.join(code_lines)})
            i += 1
            continue

        # Horizontal rules
        if re.match(r'^(-{3,}|\*{3,}|_{3,})$', stripped):
            elements.append({'type': 'hr'})
            i += 1
            continue

        # Headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if header_match:
            level = len(header_match.group(1))
            content = header_match.group(2).strip()
            elements.append({'type': 'heading', 'level': level, 'content': content})
            i += 1
            continue

        # Unordered list items
        if re.match(r'^[-*+]\s+', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\s*[-*+]\s+', lines[i].strip()):
                item = re.sub(r'^\s*[-*+]\s+', '', lines[i].strip())
                list_items.append(item)
                i += 1
            elements.append({'type': 'ul', 'items': list_items})
            continue

        # Ordered list items
        if re.match(r'^\d+\.\s+', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i].strip()):
                item = re.sub(r'^\s*\d+\.\s+', '', lines[i].strip())
                list_items.append(item)
                i += 1
            elements.append({'type': 'ol', 'items': list_items})
            continue

        # Tables - look for table pattern
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            elements.append({'type': 'table', 'rows': table_lines})
            continue

        # Regular paragraph - collect consecutive non-empty, non-special lines
        para_lines = []
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                break
            if re.match(r'^#{1,6}\s+', l) or l.startswith('```') or \
               re.match(r'^[-*+]\s+', l) or re.match(r'^\d+\.\s+', l) or \
               l.startswith('|') or re.match(r'^(-{3,}|\*{3,}|_{3,})$', l):
                break
            para_lines.append(l)
            i += 1
        if para_lines:
            elements.append({'type': 'paragraph', 'content': ' '.join(para_lines)})

    return elements


def bold_text(paragraph: object, text: str) -> None:
    """Add bold text to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = True


def italic_text(paragraph: object, text: str) -> None:
    """Add italic text to a paragraph."""
    run = paragraph.add_run(text)
    run.italic = True


def add_formatted_paragraph(doc: Document, content: str) -> None:
    """Add a paragraph with bold and inline code formatting."""
    # Split by backticks for inline code
    parts = re.split(r'`([^`]+)`', content)
    para = doc.add_paragraph()
    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            # This is inside backticks - make it code style
            run = para.add_run(f'{part}')
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # Regular text - apply bold markers
            if part.startswith('**') and part.endswith('**'):
                bold_text(para, part[2:-2])
            elif part.startswith('*') and part.endswith('*'):
                italic_text(para, part[1:-1])
            else:
                para.add_run(part)


def add_table(doc: Document, rows: list[str], num_cols: int) -> None:
    """Add a markdown table to the document."""
    table = doc.add_table(rows=num_cols + 1, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = [cell.strip() for cell in rows[0].split('|') if cell.strip()]
    if len(header_cells) > 0 and header_cells[0] == '':
        header_cells = header_cells[1:]
    for j, cell_text in enumerate(header_cells[:num_cols]):
        table.rows[0].cells[j].text = cell_text

    # Data rows (skip separator row)
    data_rows = [r for r in rows[2:] if not all(c in '-=| ' for c in r)]
    for i, row_str in enumerate(data_rows[:num_cols]):
        cells = [cell.strip() for cell in row_str.split('|') if cell.strip()]
        if len(cells) > 0 and cells[0] == '':
            cells = cells[1:]
        for j, cell_text in enumerate(cells[:num_cols]):
            table.rows[i + 1].cells[j].text = cell_text


def add_code_block(doc: Document, content: str, language: str) -> None:
    """Add a code block to the document."""
    para = doc.add_paragraph()
    run = para.add_run(f'// {language}: {content[:60]}...' if len(content) > 60 else f'// {language}: {content}')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def create_docx_from_markdown(md_path: Path, output_dir: Path, diagrams: list[dict], rendered: dict[int, Path]) -> Path:
    """Create a Word document from markdown with embedded Mermaid diagram images."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Architecture Overview: Integrated Sailing Platform (RegateVeleLiche)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_items = [
        'Executive Summary',
        'System Architecture',
        'Technology Stack',
        'Backend Services Architecture',
        'Frontend Architecture',
        'Database Schema',
        'API Endpoints Reference',
        'Algorithm Modules',
        'Key User Flows',
        'Roles and Permissions',
        'Deployment Architecture'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    # Read markdown content
    md_text = md_path.read_text(encoding='utf-8')

    # Parse elements
    elements = parse_markdown_into_elements(md_text, diagrams)

    # Add elements to document
    for elem in elements:
        if elem['type'] == 'heading':
            level_map = {1: 0, 2: 1, 3: 2}
            level = min(elem['level'], 3)
            doc.add_heading(elem['content'], level=level_map.get(level, 2))

        elif elem['type'] == 'paragraph':
            add_formatted_paragraph(doc, elem['content'])

        elif elem['type'] == 'ul':
            for item in elem['items']:
                doc.add_paragraph(item, style='List Bullet')

        elif elem['type'] == 'ol':
            for item in elem['items']:
                doc.add_paragraph(item, style='List Number')

        elif elem['type'] == 'table':
            rows = elem['rows']
            if len(rows) > 0:
                num_cols = len([c for c in rows[0].split('|') if c.strip()]) - 1
                add_table(doc, rows, min(num_cols, 6))

        elif elem['type'] == 'code':
            add_code_block(doc, elem['content'], elem.get('language', ''))

        elif elem['type'] == 'hr':
            doc.add_paragraph('_' * 50)

        elif elem['type'] == 'mermaid':
            # Find the diagram index by matching content
            for d in diagrams:
                if d['content'].strip() == elem['content'].strip():
                    img_path = rendered.get(d['index'])
                    if img_path and img_path.exists():
                        doc.add_picture(str(img_path), width=Inches(6.5))
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run('Mermaid Diagram')
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                    break

    # Save document
    output_path = output_dir / "Architecture_Overview.docx"
    doc.save(str(output_path))
    return output_path


def main():
    base_dir = Path(__file__).parent
    md_path = base_dir / "ARCHITECTURE_OVERVIEW.md"
    diagram_dir = base_dir / "diagrams"
    output_dir = base_dir

    if not md_path.exists():
        print(f"Error: {md_path} not found", file=sys.stderr)
        sys.exit(1)

    # Create diagrams directory
    diagram_dir.mkdir(exist_ok=True)

    # Extract mermaid blocks
    md_text = md_path.read_text(encoding='utf-8')
    diagrams = extract_mermaid_blocks(md_text)
    print(f"Found {len(diagrams)} Mermaid diagrams")

    # Render all diagrams to PNG
    rendered = render_all_diagrams(diagram_dir, diagrams)
    print(f"\nRendered {len(rendered)} of {len(diagrams)} diagrams successfully")

    if not rendered:
        print("No diagrams were rendered. Check for errors above.", file=sys.stderr)
        sys.exit(1)

    # Create Word document
    output_path = create_docx_from_markdown(md_path, output_dir, diagrams, rendered)
    print(f"\nDocument saved to: {output_path}")


if __name__ == '__main__':
    main()